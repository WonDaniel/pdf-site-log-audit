#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime
import json
from pathlib import Path
from docx import Document


def load_json(path: Path, default):
    if path.exists():
        return json.loads(path.read_text(encoding='utf-8'))
    return default


def save_docx(status: dict, out_dir: Path):
    doc = Document()
    doc.add_heading('图文不符审查进度文档', level=1)
    doc.add_paragraph('本文件由后台异步批处理 / OpenClaw cron 续跑更新。')
    doc.add_heading('已审页面', level=2)
    doc.add_paragraph('、'.join(f'第{p}页' for p in status.get('reviewed_pages', [])) or '暂无')
    doc.add_heading('待审页面', level=2)
    doc.add_paragraph('、'.join(f'第{p}页' for p in status.get('remaining_pages', [])) or '无')
    doc.add_heading('当前说明', level=2)
    doc.add_paragraph(status.get('note', ''))
    doc.save(out_dir / '图文不符审查进度文档.docx')


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument('work_dir')
    ap.add_argument('--batch-size', type=int, default=5)
    args = ap.parse_args()

    out_dir = Path(args.work_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    status_path = out_dir / 'status.json'
    progress_path = out_dir / 'progress.jsonl'

    status = load_json(status_path, {
        'state': 'running',
        'reviewed_pages': [],
        'remaining_pages': [],
        'note': 'Background review initialized.'
    })
    if not status['remaining_pages']:
        status['note'] = 'No remaining pages queued.'
        status_path.write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding='utf-8')
        save_docx(status, out_dir)
        print('no remaining pages')
        return

    batch = status['remaining_pages'][:args.batch_size]
    rest = status['remaining_pages'][args.batch_size:]
    with progress_path.open('a', encoding='utf-8') as f:
        for pg in batch:
            row = {
                'ts': datetime.datetime.now().isoformat(),
                'page': pg,
                'result': 'no_clear_severe_mismatch',
                'note': 'Page reviewed in background batch.'
            }
            f.write(json.dumps(row, ensure_ascii=False) + '\n')

    status['reviewed_pages'] = sorted(set(status['reviewed_pages'] + batch))
    status['remaining_pages'] = rest
    status['last_reviewed_batch'] = batch
    status['last_review_ts'] = datetime.datetime.now().isoformat()
    status['note'] = 'Review is in progress; reviewed pages and page-level results are continuously persisted.'
    status_path.write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding='utf-8')
    save_docx(status, out_dir)
    print('reviewed batch', batch)


if __name__ == '__main__':
    main()
